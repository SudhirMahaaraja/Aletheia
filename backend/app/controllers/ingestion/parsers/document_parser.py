import logging
import re

from app.controllers.ingestion.parsers.base_parser import ParsedChunk

logger = logging.getLogger(__name__)

MIN_CHUNK_LENGTH = 50


async def parse_pdf(file_bytes: bytes, filename: str) -> list[ParsedChunk]:
    import fitz  # pymupdf

    chunks: list[ParsedChunk] = []
    try:
        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages_text: list[str] = []
        for page in doc:
            text = page.get_text()
            if text and text.strip():
                pages_text.append(text)
        doc.close()
    except Exception as exc:
        logger.error("Failed to parse PDF %s: %s", filename, exc)
        return []

    # Group every 2 pages into one chunk
    for i in range(0, len(pages_text), 2):
        combined = "\n\n".join(pages_text[i:i + 2])
        if len(combined.strip()) < MIN_CHUNK_LENGTH:
            continue
        page_label = f"p{i + 1}" if i + 1 >= len(pages_text) else f"p{i + 1}-{min(i + 2, len(pages_text))}"
        chunks.append(ParsedChunk(
            chunk_type="section",
            name=f"{filename}_{page_label}",
            content=combined,
            language="text",
            file_path=filename,
            line_start=None,
            line_end=None,
            imports=[],
            calls=[],
            metadata={"page_number": i + 1},
        ))

    from app.controllers.ingestion.parsers.base_parser import post_process_chunks
    return post_process_chunks(chunks)


async def parse_docx(file_bytes: bytes, filename: str) -> list[ParsedChunk]:
    import io
    from docx import Document

    chunks: list[ParsedChunk] = []
    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        logger.error("Failed to parse DOCX %s: %s", filename, exc)
        return []

    current_heading = "Introduction"
    current_paragraphs: list[str] = []
    sections: list[tuple[str, str]] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if para.style and para.style.name and para.style.name.startswith("Heading"):
            if current_paragraphs:
                sections.append((current_heading, "\n\n".join(current_paragraphs)))
            current_heading = text
            current_paragraphs = []
        else:
            current_paragraphs.append(text)

    if current_paragraphs:
        sections.append((current_heading, "\n\n".join(current_paragraphs)))

    # If no headings were found, group every 10 paragraphs
    if len(sections) <= 1 and not any(
        p.style and p.style.name and p.style.name.startswith("Heading") for p in doc.paragraphs
    ):
        all_text = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        sections = []
        for i in range(0, len(all_text), 10):
            batch = all_text[i:i + 10]
            sections.append((f"Section_{i // 10 + 1}", "\n\n".join(batch)))

    for heading, body in sections:
        if len(body.strip()) < MIN_CHUNK_LENGTH:
            continue
        chunks.append(ParsedChunk(
            chunk_type="section",
            name=heading,
            content=body,
            language="text",
            file_path=filename,
            line_start=None,
            line_end=None,
            imports=[],
            calls=[],
            metadata={"heading": heading},
        ))

    from app.controllers.ingestion.parsers.base_parser import post_process_chunks
    return post_process_chunks(chunks)


async def parse_markdown(content: str, filename: str) -> list[ParsedChunk]:
    chunks: list[ParsedChunk] = []
    sections: list[tuple[str, list[str]]] = []
    current_heading = "Introduction"
    current_lines: list[str] = []

    for line in content.splitlines():
        if line.startswith("## "):
            if current_lines:
                sections.append((current_heading, current_lines))
            current_heading = line.lstrip("#").strip()
            current_lines = [line]
        else:
            current_lines.append(line)

    if current_lines:
        sections.append((current_heading, current_lines))

    min_len = 1

    for heading, lines in sections:
        body = "\n".join(lines)
        if len(body.strip()) < min_len:
            continue
        chunks.append(ParsedChunk(
            chunk_type="section",
            name=heading,
            content=body,
            language="markdown",
            file_path=filename,
            line_start=None,
            line_end=None,
            imports=[],
            calls=[],
            metadata={"heading": heading},
        ))

    from app.controllers.ingestion.parsers.base_parser import post_process_chunks
    return post_process_chunks(chunks)


async def parse_txt(content: str, filename: str) -> list[ParsedChunk]:
    chunks: list[ParsedChunk] = []
    words = content.split()

    is_readme = filename.lower().split("/")[-1].startswith("readme")
    min_len = 1 if is_readme else MIN_CHUNK_LENGTH

    for i in range(0, len(words), 800):
        batch = " ".join(words[i:i + 800])
        if len(batch.strip()) < min_len:
            continue
        chunk_idx = i // 800
        chunks.append(ParsedChunk(
            chunk_type="section",
            name=f"{filename}_chunk_{chunk_idx}",
            content=batch,
            language="text",
            file_path=filename,
            line_start=None,
            line_end=None,
            imports=[],
            calls=[],
            metadata={"chunk_index": chunk_idx},
        ))

    from app.controllers.ingestion.parsers.base_parser import post_process_chunks
    return post_process_chunks(chunks)
