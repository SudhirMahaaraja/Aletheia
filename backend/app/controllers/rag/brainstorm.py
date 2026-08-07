import logging
import json
from datetime import datetime, timezone
import openai
from bson import ObjectId

from app.core.config import get_settings
from app.controllers.rag.retriever import HybridRetriever

logger = logging.getLogger(__name__)


def save_as_docx(markdown_content: str, file_path: str, title: str = "Brainstorm Report"):
    from docx import Document
    from docx.shared import Inches, Pt
    import re
    
    doc = Document()
    
    # Page setup - Margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    def add_paragraph_with_formatting(text: str, style=None):
        if style:
            p = doc.add_paragraph(style=style)
        else:
            p = doc.add_paragraph()
            
        parts = re.split(r"(\*\*.*?\*\*|\*.*?\*|`.*?`)", text)
        for part in parts:
            if not part:
                continue
            if part.startswith("**") and part.endswith("**"):
                run = p.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("*") and part.endswith("*"):
                run = p.add_run(part[1:-1])
                run.italic = True
            elif part.startswith("`") and part.endswith("`"):
                run = p.add_run(part[1:-1])
                run.font.name = 'Courier New'
                run.font.size = Pt(9.5)
            else:
                p.add_run(part)
        return p

    # Simple markdown-to-docx parsing
    lines = markdown_content.split("\n")
    
    # Extract the first H1 to use as docx title
    extracted_title = None
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("# "):
            extracted_title = stripped[2:].strip().replace("*", "").replace("`", "")
            break
            
    if extracted_title:
        doc.add_heading(extracted_title, 0)
        skip_first_h1 = True
    else:
        doc.add_heading(title, 0)
        skip_first_h1 = False

    in_code_block = False
    code_text = []
    first_h1_seen = False
    in_toc = False
    
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            if in_code_block:
                # End of code block
                if code_text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.5)
                    p.paragraph_format.space_before = Pt(6)
                    p.paragraph_format.space_after = Pt(6)
                    run = p.add_run("\n".join(code_text))
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9.5)
                code_text = []
                in_code_block = False
            else:
                in_code_block = True
            continue
            
        if in_code_block:
            code_text.append(line)
            continue
            
        # TOC Section check
        if stripped.startswith("#"):
            if "table of contents" in stripped.lower():
                in_toc = True
                continue
            else:
                in_toc = False

        if in_toc:
            continue

        # Skip specific footer/meta lines
        lower_line = stripped.lower()
        if (
            "document created:" in lower_line
            or "this document can be saved as a word" in lower_line
            or "or pdf for further distribution" in lower_line
            or "further distribution and reference" in lower_line
        ):
            continue
            
        if stripped.startswith("# "):
            if skip_first_h1 and not first_h1_seen:
                first_h1_seen = True
                continue
            doc.add_heading(stripped[2:], level=1)
        elif stripped.startswith("## "):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("#### "):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith("- ") or stripped.startswith("* "):
            add_paragraph_with_formatting(stripped[2:], style='List Bullet')
        elif re.match(r"^\d+\.\s", stripped):
            match = re.match(r"^(\d+)\.\s(.*)$", stripped)
            add_paragraph_with_formatting(match.group(2), style='List Number')
        elif stripped == "":
            continue
        else:
            add_paragraph_with_formatting(line)
            
    doc.save(file_path)


def _get_openai_client():
    settings = get_settings()
    if settings.OPENAI_ENDPOINT and "azure" in settings.OPENAI_ENDPOINT.lower():
        return openai.AsyncAzureOpenAI(
            api_key=settings.OPENAI_API_KEY,
            azure_endpoint=settings.OPENAI_ENDPOINT,
            api_version=settings.OPENAI_API_VERSION or "2024-08-01-preview",
        )
    return openai.AsyncOpenAI(api_key=settings.OPENAI_API_KEY)


class BrainstormService:
    @staticmethod
    async def stream_response(
        db,
        session_id: str,
        user_id: str,
        question: str,
        chat_history: list,
        selected_repos: list[str] | None = None,
        ip_address: str = "",
    ):
        """
        Streams a RAG response for Brainstorm Mode.
        Decomposes the query, retrieves iteratively across code and documents,
        and produces a structured gap analysis and synthesis.
        """
        settings = get_settings()
        client = _get_openai_client()
        now = datetime.now(timezone.utc)

        # 1. Fetch current session context summary
        user = await db.users.find_one({"_id": ObjectId(user_id)})
        is_admin = user and user.get("role") in ("admin", "superadmin")

        query = {"_id": session_id}
        if not is_admin:
            query["user_id"] = user_id

        session_doc = await db.chat_sessions.find_one(query)
        if not session_doc and ObjectId.is_valid(session_id):
            query_oid = {"_id": ObjectId(session_id)}
            if not is_admin:
                query_oid["user_id"] = user_id
            session_doc = await db.chat_sessions.find_one(query_oid)

        context_summary = ""
        if session_doc:
            context_summary = session_doc.get("context_summary", "")
            if not selected_repos:
                selected_repos = session_doc.get("selected_repos", [])

        # Send an initial SSE update indicating research decomposition is starting
        yield f"data: {json.dumps({'status': 'Decomposing requirements and starting multi-iteration research...'})}\n\n"

        # 2. Decompose user request into 3 search queries (autoresearch loop)
        search_queries = [question]
        if settings.OPENAI_API_KEY:
            decomposition_prompt = (
                "You are an expert system architect and analyst.\n"
                "Given the user request, decompose it into 3 distinct, highly targeted search queries "
                "to find relevant source code, architectural patterns, and business logic documents in our knowledge base.\n"
                "Output ONLY the 3 search queries, one per line. Do not add numbers, bullet points, introduction, or markdown fences.\n\n"
                f"User Request: {question}"
            )
            try:
                resp = await client.chat.completions.create(
                    model=settings.OPENAI_CHAT_MODEL,
                    messages=[{"role": "user", "content": decomposition_prompt}],
                    temperature=0.2,
                )
                content = resp.choices[0].message.content.strip()
                lines = [line.strip() for line in content.split("\n") if line.strip()]
                # Clean up any bullet prefixes if LLM ignored instructions
                cleaned_lines = []
                for line in lines:
                    for prefix in ["1. ", "2. ", "3. ", "- ", "* "]:
                        if line.startswith(prefix):
                            line = line[len(prefix):]
                    cleaned_lines.append(line.strip(' "'))
                if cleaned_lines:
                    search_queries = cleaned_lines[:3]
            except Exception as e:
                logger.warning(f"Query decomposition failed: {e}")

        logger.info(f"Brainstorm decomposed queries: {search_queries}")
        search_status = f"Searching knowledge base for: {', '.join(search_queries)}..."
        yield f"data: {json.dumps({'status': search_status})}\n\n"

        # 3. Perform iterative retrieval for each decomposed query
        all_retrieved_chunks = []
        for q in search_queries:
            chunks = await HybridRetriever.retrieve(
                db=db,
                query=q,
                mode="brainstorm",
                top_k=3,
                repo_filter=selected_repos,
            )
            all_retrieved_chunks.extend(chunks)

        # 4. Deduplicate retrieved chunks and build context
        unique_chunks = {}
        for chunk in all_retrieved_chunks:
            unique_chunks[chunk["chunk_id"]] = chunk

        context_str = ""
        sources = []
        for chunk_id, chunk in unique_chunks.items():
            payload = chunk["payload"]
            collection = chunk["collection"]

            if collection == "code_chunks":
                name = payload.get("name", "Unnamed Code")
                file_path = payload.get("file_path", "")
                repo_name = payload.get("repo_name", "")
                content = payload.get("content", "")
                lang = payload.get("language", "")
                context_str += f"[Code Chunk | Repo: {repo_name} | Path: {file_path} | Name: {name}]\n"
                context_str += f"```{lang}\n{content}\n```\n"
                sources.append({
                    "type": "code",
                    "repo_name": repo_name,
                    "file_path": file_path,
                    "name": name,
                    "chunk_id": str(chunk_id) if chunk_id is not None else ""
                })
            else:
                doc_title = payload.get("document_title", "Untitled Document")
                sec_heading = payload.get("section_heading", "")
                content = payload.get("content", "")
                context_str += f"[Document Chunk | Title: {doc_title} | Section: {sec_heading}]\n"
                context_str += f"{content}\n"
                sources.append({
                    "type": "document",
                    "document_title": doc_title,
                    "section_heading": sec_heading,
                    "chunk_id": str(chunk_id) if chunk_id is not None else ""
                })

            if chunk["relationships"]:
                context_str += "Graph Relationships:\n"
                for rel in chunk["relationships"]:
                    context_str += f"- {rel}\n"
            context_str += "---\n\n"

        # Deduplicate sources
        unique_sources = []
        seen_chunks = set()
        for src in sources:
            cid = src["chunk_id"]
            if cid not in seen_chunks:
                seen_chunks.add(cid)
                unique_sources.append(src)

        yield f"data: {json.dumps({'status': 'Analyzing findings and generating structured analysis...'})}\n\n"

        # 5. Construct Synthesis and Gap Analysis Prompt
        system_prompt = (
            "You are the senior lead architect for Binary Wave Solutions.\n"
            "You are in Brainstorm Mode, synthesizing new client requirements against our internal documentation and codebase.\n"
            "Your output must be a comprehensive structured analysis containing:\n"
            "1. **Summary of Findings**: Overview of what existing systems/docs are related to the client requirement.\n"
            "2. **Gap Analysis**: Detailed breakdown of what features/modules already exist that can be reused, and what components are missing and need to be built from scratch.\n"
            "3. **Proposed Architecture/Implementation Plan**: Step-by-step technical plan showing how to build the missing features, how they fit into the existing architecture, and how to wire them to the existing codebase.\n"
            "Format your response in beautiful, clear Markdown. Be highly technical, precise, and refer directly to files and concepts from the context."
        )

        if context_summary:
            system_prompt += f"\n\nHere is a summary of the conversation history for context:\n{context_summary}"

        messages = [
            {"role": "system", "content": f"{system_prompt}\n\nSearch Context (decomposed queries findings):\n{context_str or 'No relevant code or documents found.'}"}
        ]

        for msg in chat_history[-4:]:
            messages.append({"role": msg["role"], "content": msg["content"]})

        messages.append({"role": "user", "content": question})

        # 6. Stream synthesis response
        full_answer = ""
        try:
            if not settings.OPENAI_API_KEY:
                fallback = (
                    "OpenAI API key is missing. However, here is the merged context from the multi-iteration research:\n\n"
                    f"{context_str or 'No context found.'}"
                )
                yield f"data: {json.dumps({'token': fallback})}\n\n"
                full_answer = fallback
            else:
                stream = await client.chat.completions.create(
                    model=settings.OPENAI_CHAT_MODEL,
                    messages=messages,
                    temperature=0.3,
                    stream=True,
                )
                async for chunk in stream:
                    delta = chunk.choices[0].delta if chunk.choices else None
                    if delta and delta.content:
                        token = delta.content
                        full_answer += token
                        yield f"data: {json.dumps({'token': token})}\n\n"
        except Exception as e:
            error_msg = f"Error generating answer from OpenAI LLM: {str(e)}"
            yield f"data: {json.dumps({'token': error_msg})}\n\n"
            full_answer = error_msg

        # Check if user requested a report/document
        should_create_document = any(
            kw in question.lower()
            for kw in ["report", "document", "docx", "doc", "write-up"]
        )

        if should_create_document and full_answer and not full_answer.startswith("Error"):
            try:
                from app.controllers.vault_manager import resolve_vault_dir
                
                vault_dir = resolve_vault_dir(settings.VAULT_PATH)
                generated_dir = vault_dir / "generated"
                generated_dir.mkdir(parents=True, exist_ok=True)
                
                # Extract first H1 heading from LLM response to use as the title/filename
                doc_title = "Brainstorm Report"
                for line in full_answer.split("\n"):
                    stripped_line = line.strip()
                    if stripped_line.startswith("# "):
                        doc_title = stripped_line[2:].strip().replace("*", "").replace("`", "")
                        break
                
                # Clean up title for filename: remove common prefixes like "Document for"
                title_for_filename = doc_title
                import re
                prefix_pattern = re.compile(
                    r"^(document\s+for|document\s+of|brainstorm\s+report\s*:?|brainstorm\s+analysis\s*:?)\s*",
                    re.IGNORECASE
                )
                title_for_filename = prefix_pattern.sub("", title_for_filename).strip()
                
                clean_title = "".join(c for c in title_for_filename if c.isalnum() or c in " -_")[:60].strip("-_ ").replace(" ", "_")
                if not clean_title:
                    clean_title = "Brainstorm_Report"
                
                # Use only the clean title with no timestamps or numbers behind it
                filename = f"{clean_title}.docx"
                file_path = generated_dir / filename
                
                import asyncio
                def _build_doc():
                    save_as_docx(full_answer, str(file_path), title=doc_title)
                
                await asyncio.to_thread(_build_doc)
                
                # Send SSE token info about the document
                doc_msg = f"\n\n📄 **Document Generated**: Created report as Word document at `generated/{filename}` inside the vault."
                yield f"data: {json.dumps({'token': doc_msg})}\n\n"
                full_answer += doc_msg
            except Exception as docx_exc:
                logger.error(f"Failed to create docx document: {docx_exc}")

        # Resolve node_id for each source
        for src in unique_sources:
            if src.get("type") == "document":
                src["node_id"] = src.get("chunk_id")
            else:
                node_id = None
                repo_name = src.get("repo_name", "")
                file_path = src.get("file_path", "")
                name = src.get("name", "")
                
                node = await db.graph_nodes.find_one({
                    "repo_name": repo_name,
                    "file_path": file_path,
                    "name": name
                })
                if node:
                    node_id = str(node["_id"])
                else:
                    node = await db.graph_nodes.find_one({
                        "repo_name": repo_name,
                        "file_path": file_path,
                        "type": "File"
                    })
                    if node:
                        node_id = str(node["_id"])
                    else:
                        node = await db.graph_nodes.find_one({
                            "repo_name": repo_name,
                            "type": "Repository"
                        })
                        if node:
                            node_id = str(node["_id"])
                src["node_id"] = node_id

        # Send done event with sources and generated document reference
        done_data = {"done": True, "sources": unique_sources}
        if should_create_document and "filename" in locals():
            done_data["generated_doc"] = {
                "filename": filename,
                "path": f"/api/v1/chat/download/{filename}"
            }
        yield f"data: {json.dumps(done_data)}\n\n"

        # 7. Save user and assistant messages to database
        user_msg_doc = {
            "session_id": session_id,
            "role": "user",
            "content": question,
            "sources": [],
            "created_at": now,
        }
        assistant_msg_doc = {
            "session_id": session_id,
            "role": "assistant",
            "content": full_answer,
            "sources": unique_sources,
            "created_at": datetime.now(timezone.utc),
        }
        if should_create_document and "filename" in locals():
            assistant_msg_doc["generated_doc"] = {
                "filename": filename,
                "path": f"/api/v1/chat/download/{filename}"
            }

        try:
            await db.chat_messages.insert_one(user_msg_doc)
            await db.chat_messages.insert_one(assistant_msg_doc)
        except Exception as e:
            logger.error(f"Failed to insert chat messages: {e}")

        # 8. Update message count and handle rolling summary
        msg_count = await db.chat_messages.count_documents({"session_id": session_id})
        
        new_summary = context_summary
        if msg_count >= 16 and settings.OPENAI_API_KEY:
            cursor = db.chat_messages.find({"session_id": session_id}).sort("created_at", 1)
            all_msgs = await cursor.to_list(length=100)
            
            summary_prompt = (
                "You are a conversation memory manager. Generate a concise summary of the following conversation history.\n"
                "Focus on the client requirements and technical gap analyses discussed.\n"
                "Keep it under 200 words.\n\n"
                "Conversation History:\n"
            )
            for m in all_msgs:
                role = "User" if m["role"] == "user" else "Assistant"
                summary_prompt += f"{role}: {m['content']}\n"
            summary_prompt += "\nConcise Summary:"

            try:
                resp = await client.chat.completions.create(
                    model=settings.OPENAI_CHAT_MODEL,
                    messages=[{"role": "user", "content": summary_prompt}],
                    temperature=0.2,
                )
                new_summary = resp.choices[0].message.content.strip()
                logger.info(f"Generated rolling summary for session {session_id}")
            except Exception as e:
                logger.error(f"Failed to generate rolling summary: {e}")

        # Update chat_session document
        try:
            update_data = {
                "message_count": msg_count,
                "updated_at": datetime.now(timezone.utc),
            }
            if new_summary != context_summary:
                update_data["context_summary"] = new_summary

            try:
                query_ids = [session_id]
                if ObjectId.is_valid(session_id):
                    query_ids.append(ObjectId(session_id))
                await db.chat_sessions.update_one({"_id": {"$in": query_ids}}, {"$set": update_data})
            except Exception as e:
                logger.error(f"Failed to update chat session: {e}")
        except Exception as e:
            logger.error(f"Failed to update chat session: {e}")

        # 9. Write audit log
        try:
            await db.audit_logs.insert_one({
                "user_id": user_id,
                "action": "chat_message",
                "resource_type": "session",
                "resource_id": session_id,
                "detail": f"Brainstorm query: {question[:100]}...",
                "ip_address": ip_address,
                "created_at": datetime.now(timezone.utc),
            })
        except Exception as e:
            logger.warning(f"Failed to write audit log: {e}")
